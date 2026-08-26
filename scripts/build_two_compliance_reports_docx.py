import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path("/home/ubuntu/medmethod-direct")
DELIVERABLES = Path("/home/ubuntu/report-deliverables/mmd-compliance-verification")
COMPARISONS = DELIVERABLES / "comparisons"

REPORTS = [
    (
        PROJECT / "docs/Required_Changes_Verification_Report.md",
        DELIVERABLES / "Required_Changes_Verification_Report.docx",
        "Required Changes Verification",
    ),
    (
        PROJECT / "docs/Additional_Improvements_Beyond_Requirements_Report.md",
        DELIVERABLES / "Additional_Improvements_Beyond_Requirements_Report.docx",
        "Additional Improvements Beyond Requirements",
    ),
]

PURPLE = "781F7F"
PINK = "D92C94"
LIGHT_PURPLE = "F5EFF6"
GRAY = "55515A"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(document: Document, short_title: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 25, PURPLE),
        ("Heading 1", 18, PURPLE),
        ("Heading 2", 14, PURPLE),
        ("Heading 3", 12, PINK),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = f"MedMethod Direct  |  {short_title}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(section.footer.paragraphs[0])


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"\[([0-9]+)\]", r"[\1]", text)
    text = text.replace("**", "").replace("`", "")
    return text


def add_inline_runs(paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(clean_inline(text[cursor : match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(clean_inline(token[2:-2]))
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(clean_inline(text[cursor:]))


def add_table(document: Document, lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) >= 2 and all(re.fullmatch(r":?-+:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = clean_inline(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.size = Pt(8.7)
            if row_index == 0:
                set_cell_shading(cell, PURPLE)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_PURPLE)
    document.add_paragraph()


def resolve_image(source: str) -> Path:
    return COMPARISONS / Path(source).name


def add_image(document: Document, alt: str, source: str) -> None:
    image = resolve_image(source)
    if not image.exists():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"[Image unavailable: {alt}]")
        run.italic = True
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(alt)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(PURPLE)
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image), width=Inches(5.8))


def markdown_to_docx(markdown_path: Path, output_path: Path, short_title: str) -> None:
    lines = markdown_path.read_text().splitlines()
    document = Document()
    configure_document(document, short_title)
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, " ".join(paragraph_buffer))
        paragraph_buffer = []

    index = 0
    while index < len(lines):
        raw_line = lines[index]
        hard_break = raw_line.endswith("  ")
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue

        image_match = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            flush_paragraph()
            add_image(document, image_match.group(1), image_match.group(2))
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_table(document, table_lines)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1:
                paragraph = document.add_paragraph(text, style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                accent = document.add_paragraph()
                accent.paragraph_format.space_after = Pt(10)
                run = accent.add_run("━━━━━━━━━━━━━━━━━━━━━━━━")
                run.font.color.rgb = RGBColor.from_string(PINK)
            else:
                document.add_heading(text, level=level - 1)
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            paragraph = document.add_paragraph(style="Quote")
            add_inline_runs(paragraph, stripped[1:].strip())
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(GRAY)
            index += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, stripped[2:])
            index += 1
            continue

        reference = re.fullmatch(r'\[(\d+)\]:\s+(\S+)(?:\s+"([^"]+)")?', stripped)
        if reference:
            flush_paragraph()
            number, source, description = reference.groups()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.hanging_indent = Inches(0.2)
            paragraph.paragraph_format.space_after = Pt(3)
            display = f"[{number}] {description or source} — Source: {source}"
            add_inline_runs(paragraph, display)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor.from_string(GRAY)
            index += 1
            continue

        paragraph_buffer.append(stripped)
        if hard_break:
            flush_paragraph()
        index += 1

    flush_paragraph()
    document.core_properties.author = "Manus AI"
    document.core_properties.subject = short_title
    document.core_properties.title = clean_inline(lines[0].lstrip("# "))
    document.save(output_path)
    print(output_path)


def main() -> None:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    for markdown_path, output_path, short_title in REPORTS:
        markdown_to_docx(markdown_path, output_path, short_title)


if __name__ == "__main__":
    main()
